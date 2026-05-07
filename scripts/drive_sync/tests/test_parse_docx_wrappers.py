"""Regression tests: the parser must descend through OOXML transparent wrappers.

OOXML lets editors (especially Google Docs export) wrap content in elements
that carry no semantics of their own — `<w:sdt>` (structured document tag),
`<mc:AlternateContent>` (Office compatibility), `<w:ins>` (tracked-change
insertion). If the body iterator only matches direct `<w:p>`/`<w:tbl>`
children, anything wrapped is silently dropped. These tests construct
synthetic docx files exercising each wrapper at block- and inline-level and
assert the wrapped content surfaces through `parse_docx`.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

from drive_sync.models import LinkRun, Paragraph, Table, TextRun
from drive_sync.parse.docx import ParseDocxOptions, parse_docx
from drive_sync.report import ParseReport


_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"


def _W(tag: str) -> str:
    return f"{{{_W_NS}}}{tag}"


def _MC(tag: str) -> str:
    return f"{{{_MC_NS}}}{tag}"


def _build_base_doc(tmp_path: Path) -> tuple[Document, Path]:
    """Create a minimal-but-valid info.docx with the required Metadata + one
    content section (`# Introduction`). Returns (doc, path)."""
    doc = Document()
    doc.add_heading("Metadata", level=1)
    table = doc.add_table(rows=6, cols=2)
    pairs = [
        ("Key", "Value"),
        ("title", "Test U"),
        ("sidebar_label", "Test"),
        ("sidebar_position", "99"),
        ("apply_url", "https://example.com/apply"),
        ("page_h1", "Test U"),
    ]
    for r, (k, v) in zip(table.rows, pairs):
        r.cells[0].text = k
        r.cells[1].text = v
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("Plain intro paragraph.")
    path = tmp_path / "info.docx"
    doc.save(str(path))
    return doc, path


def _xml_to_elem(xml: str):
    """Parse a fragment with both w:* and mc:* namespaces declared."""
    return etree.fromstring(
        f'<root xmlns:w="{_W_NS}" xmlns:mc="{_MC_NS}">{xml}</root>'
    )[0]


def _inject_into_body(path: Path, fragment: str) -> None:
    """Re-open `path`, insert `fragment` into the body just before <w:sectPr>,
    and save back."""
    doc = Document(str(path))
    body = doc.element.body
    new_el = _xml_to_elem(fragment)
    sect_pr = body.find(_W("sectPr"))
    if sect_pr is not None:
        sect_pr.addprevious(new_el)
    else:
        body.append(new_el)
    doc.save(str(path))


def _parse(path: Path):
    report = ParseReport()
    parsed = parse_docx(str(path), ParseDocxOptions(file_label="test/info.docx"), report)
    assert parsed is not None, [e.message for e in report.entries]
    return parsed, report


# ---------------------------------------------------------------------------
# Block-level wrappers
# ---------------------------------------------------------------------------


def test_sdt_wrapped_table_is_visible(tmp_path: Path) -> None:
    """A <w:tbl> inside <w:sdt>/<w:sdtContent> at body level must surface."""
    _, path = _build_base_doc(tmp_path)
    _inject_into_body(
        path,
        """
        <w:sdt>
          <w:sdtContent>
            <w:tbl>
              <w:tr>
                <w:tc><w:p><w:r><w:t>SDT-A</w:t></w:r></w:p></w:tc>
                <w:tc><w:p><w:r><w:t>SDT-B</w:t></w:r></w:p></w:tc>
              </w:tr>
              <w:tr>
                <w:tc><w:p><w:r><w:t>1</w:t></w:r></w:p></w:tc>
                <w:tc><w:p><w:r><w:t>2</w:t></w:r></w:p></w:tc>
              </w:tr>
            </w:tbl>
          </w:sdtContent>
        </w:sdt>
        """,
    )
    parsed, _ = _parse(path)
    blocks = parsed.sections["Introduction"]
    tables = [b for b in blocks if isinstance(b, Table)]
    assert len(tables) == 1, "SDT-wrapped table should appear in Introduction"
    header = [c.runs[0].text for c in tables[0].rows[0].cells]
    assert header == ["SDT-A", "SDT-B"]


def test_alternate_content_choice_branch_is_used(tmp_path: Path) -> None:
    """<mc:AlternateContent>: the modern <mc:Choice> branch surfaces; the
    legacy <mc:Fallback> branch is ignored."""
    _, path = _build_base_doc(tmp_path)
    _inject_into_body(
        path,
        """
        <mc:AlternateContent>
          <mc:Choice Requires="w14">
            <w:p><w:r><w:t>chosen-content</w:t></w:r></w:p>
          </mc:Choice>
          <mc:Fallback>
            <w:p><w:r><w:t>fallback-content</w:t></w:r></w:p>
          </mc:Fallback>
        </mc:AlternateContent>
        """,
    )
    parsed, _ = _parse(path)
    flat = " ".join(
        r.text
        for b in parsed.sections["Introduction"]
        if isinstance(b, Paragraph)
        for r in b.runs
        if isinstance(r, TextRun)
    )
    assert "chosen-content" in flat
    assert "fallback-content" not in flat


def test_tracked_insert_block_is_kept_delete_is_dropped(tmp_path: Path) -> None:
    """<w:ins> wrapping a paragraph keeps it; <w:del> removes it (accept-changes
    semantics so an editor with track-changes accidentally on still ships)."""
    _, path = _build_base_doc(tmp_path)
    _inject_into_body(
        path,
        """
        <w:ins w:id="1" w:author="x" w:date="2026-01-01T00:00:00Z">
          <w:p><w:r><w:t>kept-via-ins</w:t></w:r></w:p>
        </w:ins>
        <w:del w:id="2" w:author="x" w:date="2026-01-01T00:00:00Z">
          <w:p><w:r><w:t>dropped-via-del</w:t></w:r></w:p>
        </w:del>
        """,
    )
    parsed, _ = _parse(path)
    flat = " ".join(
        r.text
        for b in parsed.sections["Introduction"]
        if isinstance(b, Paragraph)
        for r in b.runs
        if isinstance(r, TextRun)
    )
    assert "kept-via-ins" in flat
    assert "dropped-via-del" not in flat


# ---------------------------------------------------------------------------
# Inline-level wrappers (inside <w:p>)
# ---------------------------------------------------------------------------


def test_inline_sdt_runs_surface(tmp_path: Path) -> None:
    """A <w:sdt> wrapping <w:r>s inside a <w:p>: the runs must reach IR."""
    _, path = _build_base_doc(tmp_path)
    _inject_into_body(
        path,
        """
        <w:p>
          <w:r><w:t xml:space="preserve">before-</w:t></w:r>
          <w:sdt>
            <w:sdtContent>
              <w:r><w:t>inline-sdt</w:t></w:r>
            </w:sdtContent>
          </w:sdt>
          <w:r><w:t xml:space="preserve">-after</w:t></w:r>
        </w:p>
        """,
    )
    parsed, _ = _parse(path)
    paragraphs = [b for b in parsed.sections["Introduction"] if isinstance(b, Paragraph)]
    flat = "".join(
        r.text for p in paragraphs for r in p.runs if isinstance(r, TextRun)
    )
    assert "before-inline-sdt-after" in flat


def test_hyperlink_inside_inline_ins_surfaces(tmp_path: Path) -> None:
    """A <w:hyperlink> nested inside <w:ins> inside <w:p> must produce a LinkRun."""
    _, path = _build_base_doc(tmp_path)
    # Hyperlinks reference w:r:id from the part rels — for this synthetic test
    # we use w:anchor (no relationship needed) so the URL becomes "#anchor".
    _inject_into_body(
        path,
        """
        <w:p>
          <w:ins w:id="3" w:author="x" w:date="2026-01-01T00:00:00Z">
            <w:hyperlink w:anchor="bookmark-1">
              <w:r><w:t>linked-text</w:t></w:r>
            </w:hyperlink>
          </w:ins>
        </w:p>
        """,
    )
    parsed, _ = _parse(path)
    paragraphs = [b for b in parsed.sections["Introduction"] if isinstance(b, Paragraph)]
    links = [
        r for p in paragraphs for r in p.runs if isinstance(r, LinkRun)
    ]
    assert any(l.text == "linked-text" and l.url == "#bookmark-1" for l in links), (
        f"got links: {[(l.text, l.url) for l in links]}"
    )


def test_inline_del_is_dropped(tmp_path: Path) -> None:
    """<w:del> inside a paragraph: its runs do NOT reach IR."""
    _, path = _build_base_doc(tmp_path)
    _inject_into_body(
        path,
        """
        <w:p>
          <w:r><w:t xml:space="preserve">keep-this </w:t></w:r>
          <w:del w:id="4" w:author="x" w:date="2026-01-01T00:00:00Z">
            <w:r><w:delText>not-this</w:delText></w:r>
          </w:del>
        </w:p>
        """,
    )
    parsed, _ = _parse(path)
    paragraphs = [b for b in parsed.sections["Introduction"] if isinstance(b, Paragraph)]
    flat = "".join(
        r.text for p in paragraphs for r in p.runs if isinstance(r, TextRun)
    )
    assert "keep-this" in flat
    assert "not-this" not in flat


# ---------------------------------------------------------------------------
# Run-level additions
# ---------------------------------------------------------------------------


def test_run_text_recognizes_cr_and_nbhyphen(tmp_path: Path) -> None:
    """<w:cr> → newline (like <w:br>); <w:noBreakHyphen> → U+2011."""
    _, path = _build_base_doc(tmp_path)
    _inject_into_body(
        path,
        """
        <w:p>
          <w:r>
            <w:t xml:space="preserve">a</w:t><w:cr/>
            <w:t xml:space="preserve">b</w:t><w:noBreakHyphen/>
            <w:t xml:space="preserve">c</w:t>
          </w:r>
        </w:p>
        """,
    )
    parsed, _ = _parse(path)
    paragraphs = [b for b in parsed.sections["Introduction"] if isinstance(b, Paragraph)]
    flat = "".join(
        r.text for p in paragraphs for r in p.runs if isinstance(r, TextRun)
    )
    assert "a\nb‑c" in flat
