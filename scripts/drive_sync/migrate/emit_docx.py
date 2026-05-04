"""IR → .docx with proper Word styles.

Goals:
- Heading 1 / 2 / 3 use the built-in Word heading styles (blue, sized,
  separator lines) — not just "big bold text".
- Tables use a built-in style (Light List Accent 1) with a colored header
  row, cell borders, and white body — readable, not a wall of plain text.
- Lists use 'List Bullet' / 'List Number' styles which have proper indent
  and clean bullet/number markers.
- Hyperlinks render in blue underline (the built-in `Hyperlink` character
  style) and are clickable in Word and Google Docs.
- Metadata table is a small 2-column table with bold keys and a soft
  header — visually distinct from content tables.

Design: spec/001-add-google-drive-backend-data/design.md §5.4.
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.table import _Cell

from drive_sync.models import (
    Block,
    Blockquote,
    Code,
    FacultyGroup,
    Heading,
    LinkRun,
    List_,
    ListItem,
    Metadata,
    Paragraph,
    RawHtml,
    Run,
    ScholarshipIR,
    Table,
    TableRow,
    TextRun,
    UniversityIR,
)


# Word's built-in style IDs. python-docx looks them up by name; if the docx
# template (default.docx) doesn't have them, an exception is raised. The
# defaults below are present in python-docx's bundled template.
_HEADING1 = "Heading 1"
_HEADING2 = "Heading 2"
_HEADING3 = "Heading 3"
_HEADING4 = "Heading 4"
_HEADING5 = "Heading 5"
_HEADING6 = "Heading 6"
_TITLE = "Title"
_SUBTITLE = "Subtitle"
_INTENSE_QUOTE = "Intense Quote"
_QUOTE = "Quote"
_LIST_BULLET = "List Bullet"
_LIST_NUMBER = "List Number"
_NORMAL = "Normal"

# The forward parser detects blockquotes by paragraph style 'Quote'. We emit
# 'Quote' for blockquotes so the round-trip works.

_TABLE_STYLE = "Light Grid Accent 1"
_METADATA_TABLE_STYLE = "Light List Accent 5"  # softer header for meta


def emit_docx(
    ir: UniversityIR | ScholarshipIR,
    out_path: Path,
) -> None:
    """Render `ir` to a .docx file at `out_path`. Caller ensures the parent dir exists."""
    doc = Document()
    _ensure_quote_style(doc)

    # Title block.
    title = doc.add_paragraph(ir.meta.page_h1 or ir.meta.title, style=_TITLE)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    sub_text = (
        f"Edit this document in Drive — the site rebuilds from it. "
        f"Slug: {ir.slug}.  Locale: {ir.locale}.  "
        f"Last (re)generated: {date.today().isoformat()}."
    )
    sub = doc.add_paragraph(sub_text, style=_SUBTITLE)
    for run in sub.runs:
        run.italic = True
        run.font.size = Pt(10)

    # Metadata section.
    doc.add_paragraph("Metadata", style=_HEADING1)
    _add_metadata_table(doc, ir.meta)
    doc.add_paragraph()

    if ir.kind == "university":
        _emit_university(doc, ir)
    else:
        _emit_scholarship(doc, ir)

    doc.save(str(out_path))


# ---------------------------------------------------------------------------
# University / scholarship section emission
# ---------------------------------------------------------------------------


def _emit_university(doc: Document, ir: UniversityIR) -> None:
    _emit_h1_section(doc, "Introduction", ir.introduction)
    # Note: faculty groups (`<MajorsTable>` JSX) are written to the xlsx, not docx.
    _emit_h1_section(doc, "Application", ir.application)
    tuition_label = (
        f"Tuition ({ir.tuition_year_label})" if ir.tuition_year_label else "Tuition"
    )
    _emit_h1_section(doc, tuition_label, ir.tuition)
    _emit_h1_section(doc, "Scholarships", ir.scholarships)
    _emit_h1_section(doc, "Requirements", ir.requirements)
    _emit_h1_section(doc, "Contacts", ir.contacts)


def _emit_scholarship(doc: Document, ir: ScholarshipIR) -> None:
    """Iterate whatever sections the IR has, in document order."""
    for section in ir.sections:
        _emit_h1_section(doc, section.heading, section.blocks)


def _emit_h1_section(doc: Document, heading: str, blocks: list[Block]) -> None:
    doc.add_paragraph(heading, style=_HEADING1)
    if not blocks:
        # Add a faint placeholder so editors know the section exists.
        p = doc.add_paragraph("(empty — fill me in)")
        for run in p.runs:
            run.italic = True
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        return
    for block in blocks:
        _emit_block(doc, block, depth_offset=0)
    doc.add_paragraph()  # extra spacing between sections


# ---------------------------------------------------------------------------
# Block-level rendering
# ---------------------------------------------------------------------------


def _emit_block(doc: Document, block: Block, *, depth_offset: int) -> None:
    if isinstance(block, Heading):
        d = max(1, min(6, block.depth + depth_offset))
        style = {
            1: _HEADING1,
            2: _HEADING2,
            3: _HEADING3,
            4: _HEADING4,
            5: _HEADING5,
            6: _HEADING6,
        }.get(d, _HEADING6)
        p = doc.add_paragraph(style=style)
        _add_runs(p, block.runs)
        return
    if isinstance(block, Paragraph):
        p = doc.add_paragraph()
        _add_runs(p, block.runs)
        return
    if isinstance(block, List_):
        style = _LIST_NUMBER if block.ordered else _LIST_BULLET
        for item in block.items:
            p = doc.add_paragraph(style=style)
            _add_runs(p, item.runs)
        return
    if isinstance(block, Table):
        _add_content_table(doc, block)
        return
    if isinstance(block, Blockquote):
        for p_block in block.paragraphs:
            p = doc.add_paragraph(style=_QUOTE)
            _add_runs(p, p_block.runs)
        return
    if isinstance(block, Code):
        # Render as a single paragraph with monospace font (no built-in
        # 'Code' style in default.docx, so we set the font directly).
        p = doc.add_paragraph()
        run = p.add_run(block.value)
        run.font.name = "Consolas"
        run.font.size = Pt(10)
        return
    if isinstance(block, RawHtml):
        # Raw HTML/JSX passthrough — emit each line as its own paragraph so
        # the forward parser's `_is_single_html_tag` detector picks it up.
        # We use a subtle accent color so the editor sees it's "different"
        # from prose, even though it's the literal source markup.
        from docx.shared import RGBColor as _RGB

        for line in block.content.splitlines() or [block.content]:
            stripped = line.strip()
            if not stripped:
                continue
            p = doc.add_paragraph()
            r = p.add_run(stripped)
            r.font.name = "Consolas"
            r.font.size = Pt(10)
            r.font.color.rgb = _RGB(0x88, 0x66, 0x00)  # subtle dark amber
        return


# ---------------------------------------------------------------------------
# Run-level rendering (text + hyperlinks + bold/italic/code)
# ---------------------------------------------------------------------------


def _add_runs(paragraph, runs: list[Run]) -> None:
    for run in runs:
        if isinstance(run, TextRun):
            r = paragraph.add_run(run.text)
            if run.bold:
                r.bold = True
            if run.italic:
                r.italic = True
            if run.code:
                r.font.name = "Consolas"
                r.font.size = Pt(10)
        elif isinstance(run, LinkRun):
            _add_hyperlink(paragraph, run.url, run.text, bold=run.bold, italic=run.italic)


def _add_hyperlink(paragraph, url: str, text: str, *, bold: bool = False, italic: bool = False) -> None:
    """Append a `<w:hyperlink>` to a paragraph using the built-in `Hyperlink`
    character style (blue, underlined). Click-through works in Word + Google Docs.
    """
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_style = OxmlElement("w:rStyle")
    r_style.set(qn("w:val"), "Hyperlink")
    r_pr.append(r_style)
    if bold:
        r_pr.append(OxmlElement("w:b"))
    if italic:
        r_pr.append(OxmlElement("w:i"))
    new_run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    text_el.set(qn("xml:space"), "preserve")
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# ---------------------------------------------------------------------------
# Tables
# ---------------------------------------------------------------------------


def _add_content_table(doc: Document, table: Table) -> None:
    if not table.rows:
        return
    n_cols = max(len(r.cells) for r in table.rows)
    if n_cols == 0:
        return
    t = doc.add_table(rows=len(table.rows), cols=n_cols)
    try:
        t.style = _TABLE_STYLE
    except KeyError:
        t.style = "Table Grid"  # fallback if the accent style is missing
    for r_idx, row in enumerate(table.rows):
        is_header = r_idx == 0
        for c_idx in range(n_cols):
            cell: _Cell = t.rows[r_idx].cells[c_idx]
            # Replace the default empty paragraph with our content.
            cell_para = cell.paragraphs[0]
            cell_para.text = ""  # clear default
            if c_idx < len(row.cells):
                _add_runs(cell_para, row.cells[c_idx].runs)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if is_header:
                for run in cell_para.runs:
                    run.bold = True


def _add_metadata_table(doc: Document, meta: Metadata) -> None:
    """A 2-column key/value table with a soft styled header.

    The forward parser reads this table to recover the page Metadata.
    """
    rows: list[tuple[str, str]] = [
        ("title", meta.title),
        ("sidebar_label", meta.sidebar_label),
        ("sidebar_position", str(meta.sidebar_position)),
    ]
    if meta.apply_url:
        rows.append(("apply_url", str(meta.apply_url)))
    if meta.page_h1:
        rows.append(("page_h1", meta.page_h1))

    t = doc.add_table(rows=len(rows) + 1, cols=2)
    try:
        t.style = _METADATA_TABLE_STYLE
    except KeyError:
        t.style = "Table Grid"

    # Header row.
    hdr = t.rows[0]
    h0 = hdr.cells[0].paragraphs[0]
    h0.text = "Key"
    for run in h0.runs:
        run.bold = True
    h1 = hdr.cells[1].paragraphs[0]
    h1.text = "Value"
    for run in h1.runs:
        run.bold = True

    for i, (key, value) in enumerate(rows, start=1):
        kp = t.rows[i].cells[0].paragraphs[0]
        kp.text = ""
        kr = kp.add_run(key)
        kr.bold = True
        vp = t.rows[i].cells[1].paragraphs[0]
        vp.text = value
        # If the value looks like a URL, render it as a hyperlink.
        if value.startswith(("http://", "https://")):
            vp.text = ""
            _add_hyperlink(vp, value, value)


# ---------------------------------------------------------------------------
# Style support
# ---------------------------------------------------------------------------


def _ensure_quote_style(doc: Document) -> None:
    """Make sure the 'Quote' style exists with a recognizable look.

    python-docx's default template includes 'Quote' / 'Intense Quote' but
    they vary across Word/Google Docs. We tweak it to be visually italic +
    indented + accent-colored — and matchable by the forward parser via
    `style-name='Quote'`.
    """
    styles = doc.styles
    if "Quote" not in [s.name for s in styles]:
        return
    quote_style = styles["Quote"]
    quote_style.font.italic = True
    quote_style.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


# Suppress unused-import linter noise.
_ = (FacultyGroup, ListItem, TableRow, WD_STYLE_TYPE)
