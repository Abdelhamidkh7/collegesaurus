"""docx → structured IR fragment.

Walks the document body in order via the underlying lxml elements (so paragraphs
and tables interleave correctly), splits by H1 boundaries to identify sections,
and converts each section's content into a list of `Block` objects.

Heading levels:
- The first H1 must be `Metadata` (a 2-column key/value table feeds Metadata).
- Subsequent H1s are content sections (Introduction, Application, ... or
  Overview, Grade & background requirements, ... for scholarships).
- H2s and below within a section become structural sub-headings.

Design: spec/001-add-google-drive-backend-data/design.md §3.2, §3.3, §4.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from typing import Iterator

from docx import Document as load_document
from docx.document import Document
from docx.oxml.ns import qn
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph

from drive_sync.models import (
    Block,
    Blockquote,
    Code,
    Heading,
    LinkRun,
    List_,
    ListItem,
    Metadata,
    Paragraph,
    RawHtml,
    Run,
    Table,
    TableCell,
    TableRow,
    TextRun,
)
from drive_sync.report import ParseReport


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


@dataclass
class ParsedDocx:
    """The output of `parse_docx` — sections in document order plus metadata."""

    metadata_raw: dict[str, str]
    """Raw key/value strings parsed from the # Metadata H1's table.

    Validation against `Metadata` happens later (in assemble.py), so we keep
    them as strings here.
    """

    sections: dict[str, list[Block]] = field(default_factory=dict)
    """H1 section name (verbatim, trimmed) → ordered list of Blocks.
    The Metadata section is excluded.
    """

    section_order: list[str] = field(default_factory=list)
    """H1 section names in document order (excluding Metadata)."""


@dataclass
class ParseDocxOptions:
    file_label: str
    web_view_link: str | None = None


def parse_docx(
    file_path: str,
    options: ParseDocxOptions,
    report: ParseReport,
) -> ParsedDocx | None:
    """Parse a .docx file.

    Returns None on fatal errors (which are recorded in `report`); call
    `report.has_errors()` after parsing all files to decide whether to fail
    the build.
    """
    try:
        doc = load_document(file_path)
    except Exception as err:  # noqa: BLE001 — convert any docx open error to a report entry
        report.error(
            options.file_label,
            f"cannot open docx: {err}",
            web_view_link=options.web_view_link,
        )
        return None

    sections_buckets: dict[str, list[Block]] = {}
    section_order: list[str] = []
    current: str | None = None
    current_blocks: list[Block] = []

    # We need ordered, interleaved iteration of paragraphs and tables. We also
    # need to group consecutive "List Paragraph"-styled paragraphs into a
    # single List_ block.
    pending_list: list[ListItem] | None = None
    pending_list_ordered: bool = False

    def flush_list() -> None:
        """Emit the pending list (if any) as a Block."""
        nonlocal pending_list, pending_list_ordered
        if pending_list is None:
            return
        current_blocks.append(List_(ordered=pending_list_ordered, items=pending_list))
        pending_list = None

    for item in _iter_block_items(doc):
        if isinstance(item, DocxParagraph):
            depth = _heading_depth(item)
            if depth == 1:
                # Section boundary.
                flush_list()
                if current is not None:
                    sections_buckets[current] = current_blocks
                current = _paragraph_text(item).strip()
                section_order.append(current)
                current_blocks = []
                continue

            if current is None:
                # Pre-amble before the first H1; ignore.
                continue

            if _is_list_paragraph(item):
                # Group with the pending list (or start one).
                ordered = _is_ordered_list(item)
                if pending_list is None or ordered != pending_list_ordered:
                    flush_list()
                    pending_list = []
                    pending_list_ordered = ordered
                pending_list.append(ListItem(runs=_paragraph_runs(item)))
                continue

            flush_list()

            if depth and depth >= 2:
                current_blocks.append(Heading(depth=depth, runs=_paragraph_runs(item)))
                continue

            if _is_blockquote(item):
                # Coalesce consecutive blockquoted paragraphs into one block.
                last = current_blocks[-1] if current_blocks else None
                if isinstance(last, Blockquote):
                    last.paragraphs.append(Paragraph(runs=_paragraph_runs(item)))
                else:
                    current_blocks.append(
                        Blockquote(paragraphs=[Paragraph(runs=_paragraph_runs(item))])
                    )
                continue

            text = _paragraph_text(item).strip()
            runs = _paragraph_runs(item)
            if not runs and not text:
                continue  # skip empty paragraphs

            # Raw HTML / JSX passthrough: a paragraph whose entire content is
            # a single open or close tag (e.g. `<div className="alert-warning">`
            # or `</div>`) is preserved as-is in the emitted MDX. Editors author
            # these by typing the literal text in the docx.
            if _is_single_html_tag(text):
                current_blocks.append(RawHtml(content=text))
                continue

            current_blocks.append(Paragraph(runs=runs))

        elif isinstance(item, DocxTable):
            if current is None:
                continue
            flush_list()
            current_blocks.append(_docx_table_to_block(item))

    # Flush the final section.
    flush_list()
    if current is not None:
        sections_buckets[current] = current_blocks

    # ---- Metadata extraction ----
    metadata_key = next((s for s in section_order if _normalize(s) == "metadata"), None)
    if metadata_key is None:
        report.error(
            options.file_label,
            "missing required `# Metadata` H1 at top of document",
            web_view_link=options.web_view_link,
        )
        return None

    metadata_blocks = sections_buckets[metadata_key]
    metadata_raw = _extract_metadata_table(metadata_blocks)
    if not metadata_raw:
        report.error(
            options.file_label,
            "Metadata section is missing the required 2-column key/value table",
            web_view_link=options.web_view_link,
        )
        return None

    # Build the public dict (excluding Metadata).
    public_sections: dict[str, list[Block]] = {}
    public_order: list[str] = []
    for name in section_order:
        if _normalize(name) == "metadata":
            continue
        public_sections[name] = sections_buckets[name]
        public_order.append(name)

    return ParsedDocx(metadata_raw=metadata_raw, sections=public_sections, section_order=public_order)


def find_section(parsed: ParsedDocx, name: str) -> tuple[str, list[Block]] | None:
    """Match a wanted section name against `parsed.sections` keys.

    Tolerates prefix matches (so wanted="Tuition" matches the key
    "Tuition (AY 2025-2026)"). Returns (matched_key, blocks) or None.
    """
    wanted = _normalize(name)
    for key in parsed.section_order:
        norm = _normalize(key)
        if norm == wanted or norm.startswith(wanted):
            return key, parsed.sections[key]
    return None


# ---------------------------------------------------------------------------
# Validation glue (T-2.7)
# ---------------------------------------------------------------------------


def validate_metadata(
    raw: dict[str, str],
    file_label: str,
    report: ParseReport,
    web_view_link: str | None = None,
) -> Metadata | None:
    """Validate raw Metadata table → Metadata pydantic model. Reports errors."""
    from pydantic import ValidationError

    try:
        return Metadata.model_validate(raw)
    except ValidationError as err:
        # Each error in err.errors() has loc=(field,...), msg, type
        issues = "; ".join(
            f"{'.'.join(str(p) for p in e['loc']) or '(root)'}: {e['msg']}"
            for e in err.errors()
        )
        report.error(
            file_label,
            f"Metadata validation failed — {issues}",
            web_view_link=web_view_link,
            where="Metadata table",
        )
        return None


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------


# OOXML "transparent wrapper" elements that semantically contribute no block
# of their own — their <w:p>/<w:tbl> children should bubble up as if they were
# direct body children. Without descent we silently drop content:
#   * <w:sdt>/<w:sdtContent>: structured document tag (Google Docs export
#     wraps tables in these for some content controls)
#   * <w:ins>, <w:moveTo>: tracked-change *additions* / move destinations —
#     "accept changes" semantics: keep the new content
#   * <mc:AlternateContent>/<mc:Choice>: Office compatibility wrapper — modern
#     branch
# And containers we skip entirely (their content is conceptually deleted):
#   * <w:del>, <w:moveFrom>: tracked-change *deletions* / move origins
#   * <mc:Fallback>: legacy fallback branch (only used if the Choice branch
#     is unrenderable, which never applies for our reader)

_MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"
_BLOCK_DESCEND = frozenset(
    [
        qn("w:sdt"),
        qn("w:sdtContent"),
        qn("w:ins"),
        qn("w:moveTo"),
        f"{{{_MC_NS}}}AlternateContent",
        f"{{{_MC_NS}}}Choice",
    ]
)
_BLOCK_SKIP = frozenset(
    [
        qn("w:del"),
        qn("w:moveFrom"),
        f"{{{_MC_NS}}}Fallback",
    ]
)
# Inline equivalents (for descent inside <w:p>, alongside <w:r>/<w:hyperlink>).
_INLINE_DESCEND = frozenset(
    [
        qn("w:sdt"),
        qn("w:sdtContent"),
        qn("w:ins"),
        qn("w:moveTo"),
        qn("w:smartTag"),
        qn("w:customXml"),
        qn("w:fldSimple"),
        f"{{{_MC_NS}}}AlternateContent",
        f"{{{_MC_NS}}}Choice",
    ]
)
_INLINE_SKIP = frozenset(
    [
        qn("w:del"),
        qn("w:moveFrom"),
        f"{{{_MC_NS}}}Fallback",
    ]
)


def _iter_block_items(doc: Document) -> Iterator[DocxParagraph | DocxTable]:
    """Yield paragraphs and tables in document order.

    python-docx exposes `doc.paragraphs` and `doc.tables` as flat lists each,
    losing the interleaving. We walk the body XML directly and descend through
    OOXML transparent wrappers (SDTs, tracked-change inserts, MC compat) so
    their inner <w:p>/<w:tbl> children surface to the caller.
    """
    yield from _iter_blocks(doc.element.body, doc)


def _iter_blocks(parent, doc: Document) -> Iterator[DocxParagraph | DocxTable]:
    """Recursive worker for `_iter_block_items`.

    Yields the leaf <w:p>/<w:tbl> blocks; descends into transparent wrappers;
    skips tracked-deletion wrappers (whose content is conceptually removed)
    and MC fallback branches.
    """
    for child in parent.iterchildren():
        tag = child.tag
        if tag == qn("w:p"):
            yield DocxParagraph(child, doc)
        elif tag == qn("w:tbl"):
            yield DocxTable(child, doc)
        elif tag in _BLOCK_DESCEND:
            yield from _iter_blocks(child, doc)
        # Anything else (sectPr, bookmarkStart/End, proofErr, _BLOCK_SKIP, …)
        # is correctly silent: it carries no publishable block content.


def _heading_depth(p: DocxParagraph) -> int:
    """Return 1-6 if the paragraph is styled `Heading N`, else 0."""
    style = p.style.name if p.style else ""
    if not style.startswith("Heading "):
        return 0
    try:
        depth = int(style.split()[-1])
    except (ValueError, IndexError):
        return 0
    return depth if 1 <= depth <= 6 else 0


def _paragraph_text(p: DocxParagraph) -> str:
    """Plain text of a paragraph (no formatting)."""
    return p.text


def _paragraph_runs(p: DocxParagraph) -> list[Run]:
    """Convert a paragraph's runs and hyperlinks to model Run objects, in order.

    Walks the `<w:p>` element's children in order, descending through inline
    transparent wrappers (SDT, ins, smartTag, customXml, fldSimple, MC compat)
    so runs/hyperlinks nested inside them surface as if they were direct
    children. `<w:del>` and `<w:moveFrom>` subtrees are skipped (deleted
    content).
    """
    rels = p.part.rels  # WordRelationships dict for resolving hyperlink targets
    out = list(_iter_inline_runs(p._p, rels))
    return _coalesce_runs(out)


def _iter_inline_runs(elem, rels) -> Iterator[Run]:
    """Yield Run objects from a paragraph (or transparent inline wrapper)."""
    for child in elem.iterchildren():
        tag = child.tag
        if tag == qn("w:r"):
            text = _run_text(child)
            if not text:
                continue
            bold, italic = _run_styles(child)
            yield TextRun(text=text, bold=bold, italic=italic)
        elif tag == qn("w:hyperlink"):
            url = _hyperlink_url(child, rels)
            # Hyperlinks may contain transparent wrappers too — iter() reaches
            # all <w:r> descendants regardless of wrapping.
            text = "".join(_run_text(r) for r in child.iter(qn("w:r")))
            if not text:
                continue
            bold = italic = False
            first_r = next(child.iter(qn("w:r")), None)
            if first_r is not None:
                bold, italic = _run_styles(first_r)
            yield LinkRun(text=text, url=url, bold=bold, italic=italic)
        elif tag in _INLINE_DESCEND:
            yield from _iter_inline_runs(child, rels)
        # _INLINE_SKIP and unrecognized inline siblings (bookmarkStart/End,
        # proofErr, permStart/End, …) carry no run-level text.


def _run_text(r) -> str:
    """Extract text from a `<w:r>` element (concatenating all `<w:t>` children)."""
    parts: list[str] = []
    for t in r.iterchildren():
        tag = t.tag
        if tag == qn("w:t"):
            parts.append(t.text or "")
        elif tag == qn("w:tab"):
            parts.append("\t")
        elif tag in (qn("w:br"), qn("w:cr")):
            parts.append("\n")
        elif tag == qn("w:noBreakHyphen"):
            parts.append("‑")  # U+2011, semantically a non-breaking hyphen
    return "".join(parts)


def _run_styles(r) -> tuple[bool, bool]:
    """Return (bold, italic) flags from a `<w:r>` element's `<w:rPr>`."""
    rpr = r.find(qn("w:rPr"))
    if rpr is None:
        return False, False
    bold = rpr.find(qn("w:b")) is not None
    italic = rpr.find(qn("w:i")) is not None
    return bold, italic


def _hyperlink_url(hyperlink, rels) -> str:
    """Resolve a `<w:hyperlink>` element's target URL via the part's relationships."""
    rid = hyperlink.get(qn("r:id"))
    if rid and rid in rels:
        return rels[rid].target_ref
    # Fallback: anchor (in-document link)
    anchor = hyperlink.get(qn("w:anchor"))
    return f"#{anchor}" if anchor else ""


def _coalesce_runs(runs: list[Run]) -> list[Run]:
    """Merge adjacent TextRuns with identical formatting."""
    out: list[Run] = []
    for run in runs:
        if (
            isinstance(run, TextRun)
            and out
            and isinstance(out[-1], TextRun)
            and out[-1].bold == run.bold
            and out[-1].italic == run.italic
            and out[-1].code == run.code
        ):
            out[-1] = TextRun(
                text=out[-1].text + run.text,
                bold=run.bold,
                italic=run.italic,
                code=run.code,
            )
        else:
            out.append(run)
    return out


def _is_blockquote(p: DocxParagraph) -> bool:
    style = p.style.name if p.style else ""
    return style in ("Quote", "Intense Quote")


_LIST_STYLE_PREFIXES = ("List Bullet", "List Number", "List Continue")


def _is_list_paragraph(p: DocxParagraph) -> bool:
    """A paragraph is a list item if it carries any of Word's built-in list
    styles (`List Paragraph`, `List Bullet`, `List Bullet 2`, `List Number`,
    `List Continue`, …) OR has an inline `<w:numPr>` override.

    Note: the bootstrap emit (`migrate/emit_docx.py`) writes `List Bullet` /
    `List Number` paragraphs whose bullet/number markers come from the *style*
    rather than from a per-paragraph `<w:numPr>` — so a style-name check is
    required; checking only `numPr` would miss them and the round-trip would
    flatten the list to plain paragraphs.
    """
    style = p.style.name if p.style else ""
    if style == "List Paragraph" or style.startswith(_LIST_STYLE_PREFIXES):
        return True
    p_pr = p._p.find(qn("w:pPr"))
    if p_pr is None:
        return False
    return p_pr.find(qn("w:numPr")) is not None


def _is_ordered_list(p: DocxParagraph) -> bool:
    """Determine whether a list paragraph is ordered (numbered) or unordered (bulleted).

    Two paths:
    1. **Style-driven** (built-in `List Bullet` / `List Number` family): the
       style name itself disambiguates — `List Number*` is ordered, anything
       else (`List Bullet*`, `List Continue*`, plain `List Paragraph`) is
       unordered. The bootstrap emit takes this path.
    2. **numPr-driven**: resolve `numPr/numId` → `numbering.part` and look at
       the abstract format (decimal/lowerLetter/lowerRoman → ordered;
       bullet → unordered). Used when an editor authored a list manually in
       Word/Google Docs without a list style.
    """
    style = p.style.name if p.style else ""
    if style.startswith("List Number"):
        return True
    if style.startswith(("List Bullet", "List Continue")):
        return False
    p_pr = p._p.find(qn("w:pPr"))
    if p_pr is None:
        return False
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        return False
    num_id_el = num_pr.find(qn("w:numId"))
    if num_id_el is None:
        return False
    num_id = num_id_el.get(qn("w:val"))
    # Walk numbering.xml to find the format. Conservative fallback: if we
    # can't resolve, assume bulleted.
    try:
        numbering = p.part.numbering_part.element
    except (AttributeError, KeyError):
        return False
    num_el = next(
        (n for n in numbering.iter(qn("w:num")) if n.get(qn("w:numId")) == num_id),
        None,
    )
    if num_el is None:
        return False
    abstract_num_id_el = num_el.find(qn("w:abstractNumId"))
    if abstract_num_id_el is None:
        return False
    abstract_num_id = abstract_num_id_el.get(qn("w:val"))
    abstract_num = next(
        (
            a
            for a in numbering.iter(qn("w:abstractNum"))
            if a.get(qn("w:abstractNumId")) == abstract_num_id
        ),
        None,
    )
    if abstract_num is None:
        return False
    # Check the level-0 numFmt
    lvl0 = next(
        (l for l in abstract_num.iter(qn("w:lvl")) if l.get(qn("w:ilvl")) == "0"),
        None,
    )
    if lvl0 is None:
        return False
    num_fmt = lvl0.find(qn("w:numFmt"))
    if num_fmt is None:
        return False
    fmt = num_fmt.get(qn("w:val"))
    return fmt not in ("bullet", None)


def _docx_table_to_block(t: DocxTable) -> Table:
    rows: list[TableRow] = []
    for row in t.rows:
        cells: list[TableCell] = []
        for cell in row.cells:
            # A docx table cell can contain multiple paragraphs; we flatten them.
            runs: list[Run] = []
            for para in cell.paragraphs:
                if runs:
                    runs.append(TextRun(text=" "))  # space-separate paragraphs in a cell
                runs.extend(_paragraph_runs(para))
            cells.append(TableCell(runs=runs))
        rows.append(TableRow(cells=cells))
    return Table(rows=rows)


def _extract_metadata_table(blocks: list[Block]) -> dict[str, str]:
    """The Metadata section's first Table is the key/value table.

    Skip a `Key | Value` header row if present; otherwise treat all rows as data.
    """
    table = next((b for b in blocks if isinstance(b, Table)), None)
    if table is None:
        return {}
    out: dict[str, str] = {}
    for i, row in enumerate(table.rows):
        if len(row.cells) < 2:
            continue
        key = _runs_to_text(row.cells[0].runs).strip()
        value = _runs_to_text(row.cells[1].runs).strip()
        if i == 0 and key.lower() == "key":
            continue  # skip header row
        if not key:
            continue
        out[key] = value
    return out


def _runs_to_text(runs: list[Run]) -> str:
    """Flatten runs back to plain text (formatting dropped)."""
    parts: list[str] = []
    for r in runs:
        if isinstance(r, TextRun):
            parts.append(r.text)
        elif isinstance(r, LinkRun):
            parts.append(r.text)
    return "".join(parts)


def _normalize(s: str) -> str:
    """Lower-case + strip non-alphanumerics. Preserves Unicode letters/digits.

    Used for fuzzy section name matching (e.g. `الكلية (Faculty)` and
    `Faculty` both normalize to a string containing `faculty`).
    """
    import re as _re

    return _re.sub(r"[^\w]", "", s, flags=_re.UNICODE).lower()


# A single HTML/JSX tag (open, close, or self-closing) with optional attrs.
# Used to detect "this paragraph is a passthrough markup marker like
# `<div className="alert-warning">` or `</div>`."
import re as _html_re

_SINGLE_HTML_TAG_RE = _html_re.compile(
    r"^\s*</?[A-Za-z][A-Za-z0-9_-]*(\s+[^>]*)?/?>\s*$"
)


def _is_single_html_tag(text: str) -> bool:
    return bool(_SINGLE_HTML_TAG_RE.match(text))


# Suppress "imported but unused" noise from Code being part of Block union.
_ = Code
